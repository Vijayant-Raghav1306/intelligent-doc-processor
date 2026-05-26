import textwrap

from docx import Document
from docx.document import Document as DocxDocument
from PIL import Image, ImageDraw

from app.core.logging_config import get_logger
from app.services.base_loader import BaseLoader

logger = get_logger(__name__)

# Preview canvas dimensions (pixels)
PREVIEW_WIDTH = 800
PREVIEW_HEIGHT = 1100
PREVIEW_MARGIN = 40
PREVIEW_LINE_HEIGHT = 18


class DocxLoader(BaseLoader):

    def __init__(self, file_path):
        super().__init__(file_path)  # validates file exists via BaseLoader

    # â”€â”€ Internal helpers â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

    def _open_doc(self) -> DocxDocument:
        return Document(str(self.file_path))

    def _extract_paragraph_text(self, doc: DocxDocument) -> list[str]:
        """
        Return non-empty paragraph strings.
        Filters blank paragraphs â€” DOCX uses them as visual spacing, they're noise.
        """
        return [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    def _extract_table_text(self, doc: DocxDocument) -> list[str]:
        """
        Return each table row as a ' | ' delimited string.
        Iterates doc.tables separately â€” table content does NOT appear in doc.paragraphs.
        """
        lines = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
            lines.append("")  # blank line between tables for readability
        return lines

    # â”€â”€ BaseLoader contract â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

    def extract_text(self) -> str:
        doc = self._open_doc()
        logger.info(
            "docx_parsed",
            extra={
                "doc": self.file_path.name,
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
            },
        )

        paragraph_lines = self._extract_paragraph_text(doc)
        table_lines = self._extract_table_text(doc)

        all_lines = paragraph_lines
        if table_lines:
            all_lines += ["\n--- Tables ---"] + table_lines

        text = "\n".join(all_lines)
        logger.info(
            "text_extracted",
            extra={"doc": self.file_path.name, "chars": len(text)},
        )
        return text

    def get_preview_image(self) -> Image.Image:
        """
        Render the first ~60 lines of extracted text onto a white canvas.
        DOCX has no pixel representation â€” this is the honest lightweight alternative
        to a full LibreOffice conversion.
        """
        text = self.extract_text()

        # Wrap long lines to fit within the canvas width (~95 chars at default font)
        wrapped_lines = []
        for line in text.splitlines():
            if line.strip():
                wrapped_lines.extend(textwrap.wrap(line, width=95))
            else:
                wrapped_lines.append("")  # preserve intentional blank lines

        # Cap at how many lines actually fit on the canvas
        max_lines = (PREVIEW_HEIGHT - PREVIEW_MARGIN * 2) // PREVIEW_LINE_HEIGHT
        visible_lines = wrapped_lines[:max_lines]

        # Draw onto white canvas
        img = Image.new("RGB", (PREVIEW_WIDTH, PREVIEW_HEIGHT), color=(255, 255, 255))
        draw = ImageDraw.Draw(img)

        y = PREVIEW_MARGIN
        for line in visible_lines:
            draw.text((PREVIEW_MARGIN, y), line, fill=(30, 30, 30))
            y += PREVIEW_LINE_HEIGHT

        # If content was truncated, indicate it
        if len(wrapped_lines) > max_lines:
            draw.text(
                (PREVIEW_MARGIN, PREVIEW_HEIGHT - PREVIEW_MARGIN),
                f"... ({len(wrapped_lines) - max_lines} more lines)",
                fill=(150, 150, 150),
            )

        return img

    def get_metadata(self) -> dict:
        doc = self._open_doc()
        core_props = doc.core_properties  # author, title, created, modified, etc.

        return {
            "page_count": 1,          # python-docx has no page count â€” DOCX is flow-based
            "file_type": "docx",
            "file_size_bytes": self.get_file_size_bytes(),
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "title": core_props.title or "",
            "author": core_props.author or "",
            "last_modified_by": core_props.last_modified_by or "",
        }

