"""
Shared fixtures for all tests.
pytest discovers conftest.py automatically — no imports needed in test files.

All file fixtures use tmp_path (built-in pytest fixture) so each test
gets a fresh isolated directory that is deleted after the test completes.
"""
import shutil
from pathlib import Path

import fitz  # PyMuPDF
import pytest
from docx import Document
from fastapi.testclient import TestClient
from PIL import Image, ImageDraw

from app.core.config import settings
from app.main import app

AES_256 = getattr(fitz, "PDF_ENCRYPT_AES_256", 6)

# ── Tesseract availability marker ─────────────────────────────────────────────

def _tesseract_available() -> bool:
    """Check if the Tesseract binary exists at the configured path or on PATH."""
    configured = Path(settings.tesseract_cmd)
    return configured.exists() or shutil.which("tesseract") is not None

requires_tesseract = pytest.mark.skipif(
    not _tesseract_available(),
    reason="Tesseract OCR binary not found — install it to run OCR tests",
)


# ── HTTP client ───────────────────────────────────────────────────────────────

@pytest.fixture(scope="session")
def client():
    """
    A TestClient wrapping the FastAPI app.
    scope="session" means one client for the entire test run — faster,
    since the app lifespan (directory creation, logging setup) runs once.
    """
    with TestClient(app) as c:
        yield c


# ── Test file factories ───────────────────────────────────────────────────────

@pytest.fixture
def tmp_png(tmp_path: Path) -> Path:
    """
    A real PNG file with large black text on white background.
    Text is oversized (40px) so Tesseract has a reliable chance to read it.
    Tests that assert on OCR output should keep expectations loose (not exact),
    since OCR accuracy varies by environment and Tesseract version.
    """
    img = Image.new("RGB", (600, 120), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.text((20, 30), "HELLO PYTEST", fill=(0, 0, 0))
    path = tmp_path / "test.png"
    img.save(path)
    return path


@pytest.fixture
def tmp_pdf(tmp_path: Path) -> Path:
    """
    A real single-page PDF with embedded native text (not scanned).
    Created with PyMuPDF so pdfplumber can extract it without OCR.
    """
    path = tmp_path / "test.pdf"
    doc: fitz.Document = fitz.open()
    page = doc.new_page(width=595, height=842)  # type: ignore[attr-defined]
    page.insert_text((72, 72), "Hello from pytest PDF", fontsize=16)  # type: ignore[attr-defined]
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def tmp_pdf_encrypted(tmp_path: Path) -> tuple[Path, str]:
    """
    A password-protected PDF.
    Returns (path, password) so tests can use both.
    """
    path = tmp_path / "locked.pdf"
    password = "s3cr3t"
    doc: fitz.Document = fitz.open()
    page = doc.new_page()  # type: ignore[attr-defined]
    page.insert_text((72, 72), "Secret content", fontsize=16)  # type: ignore[attr-defined]
    doc.save(
        str(path),
        encryption=AES_256,
        user_pw=password,
        owner_pw=password,
    )
    doc.close()
    return path, password


@pytest.fixture
def tmp_docx(tmp_path: Path) -> Path:
    """
    A DOCX file with both paragraph text and a table.
    Tests for table extraction must use this fixture, not tmp_png/tmp_pdf.
    """
    path = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Hello from pytest paragraph")
    doc.add_paragraph("Second paragraph for testing")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Score"
    table.cell(1, 0).text = "pytest"
    table.cell(1, 1).text = "100"

    doc.save(str(path))
    return path


@pytest.fixture
def tmp_unsupported_file(tmp_path: Path) -> Path:
    """
    A file with EXE magic bytes — guaranteed to be rejected by our MIME guard.
    We write real magic bytes (MZ header) so filetype.guess() detects it correctly.
    """
    path = tmp_path / "malicious.exe"
    path.write_bytes(b"MZ" + b"\x00" * 100)
    return path




