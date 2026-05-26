"""
Unit tests for DocxLoader.
These tests instantiate the loader directly — no HTTP, no FastAPI.
"""
import pytest
from PIL import Image

from app.services.docx_loader import DocxLoader


class TestDocxLoaderContract:
    """Verify the BaseLoader contract is fulfilled correctly."""

    def test_extract_text_returns_string(self, tmp_docx):
        text = DocxLoader(tmp_docx).extract_text()
        assert isinstance(text, str)

    def test_get_preview_image_returns_pil_image(self, tmp_docx):
        img = DocxLoader(tmp_docx).get_preview_image()
        assert isinstance(img, Image.Image)

    def test_get_metadata_returns_dict(self, tmp_docx):
        meta = DocxLoader(tmp_docx).get_metadata()
        assert isinstance(meta, dict)

    def test_missing_file_raises_file_not_found(self):
        with pytest.raises(FileNotFoundError):
            DocxLoader("does_not_exist.docx")


class TestDocxTextExtraction:
    """Verify paragraph and table content is extracted correctly."""

    def test_paragraph_text_present(self, tmp_docx):
        text = DocxLoader(tmp_docx).extract_text()
        assert "Hello from pytest paragraph" in text
        assert "Second paragraph for testing" in text

    def test_table_content_present(self, tmp_docx):
        # Table content does NOT appear in doc.paragraphs —
        # this test would fail if _extract_table_text() was missing
        text = DocxLoader(tmp_docx).extract_text()
        assert "pytest" in text
        assert "100" in text

    def test_table_cells_separated_by_pipe(self, tmp_docx):
        text = DocxLoader(tmp_docx).extract_text()
        # Our formatter joins cells with " | "
        assert "|" in text

    def test_empty_paragraphs_filtered(self, tmp_path):
        # DOCX uses blank paragraphs as spacing — they should not appear in output
        from docx import Document
        path = tmp_path / "sparse.docx"
        doc = Document()
        doc.add_paragraph("Line one")
        doc.add_paragraph("")   # blank spacing paragraph
        doc.add_paragraph("")
        doc.add_paragraph("Line two")
        doc.save(str(path))

        text = DocxLoader(path).extract_text()
        lines = [ln for ln in text.splitlines() if ln.strip()]
        assert lines == ["Line one", "Line two"]


class TestDocxMetadata:
    """Verify metadata keys and types."""

    def test_required_keys_present(self, tmp_docx):
        meta = DocxLoader(tmp_docx).get_metadata()
        for key in ("page_count", "file_type", "file_size_bytes", "paragraph_count", "table_count"):
            assert key in meta, f"Missing key: {key}"

    def test_file_type_is_docx(self, tmp_docx):
        assert DocxLoader(tmp_docx).get_metadata()["file_type"] == "docx"

    def test_table_count_correct(self, tmp_docx):
        # conftest creates one table
        assert DocxLoader(tmp_docx).get_metadata()["table_count"] == 1

    def test_file_size_bytes_positive(self, tmp_docx):
        assert DocxLoader(tmp_docx).get_metadata()["file_size_bytes"] > 0


class TestDocxPreview:
    """Verify the text-based preview image properties."""

    def test_preview_is_rgb(self, tmp_docx):
        img = DocxLoader(tmp_docx).get_preview_image()
        assert img.mode == "RGB"

    def test_preview_dimensions(self, tmp_docx):
        from app.services.docx_loader import PREVIEW_HEIGHT, PREVIEW_WIDTH
        img = DocxLoader(tmp_docx).get_preview_image()
        assert img.size == (PREVIEW_WIDTH, PREVIEW_HEIGHT)